"""Build a workout program DOCX from a workout plan JSON file."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from plan_model import default_output_paths, iter_exercise_rows, list_value, load_plan, source_warnings


TABLE_COLUMNS = [
    ("week", "Week"),
    ("day", "Day"),
    ("block", "Block"),
    ("exercise", "Exercise"),
    ("visual_source", "Visual/source"),
    ("sets_reps", "Sets x reps/time"),
    ("intensity", "Intensity"),
    ("rest", "Rest"),
    ("notes", "Notes"),
]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)

    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_text(cell, text: Any, font_size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    run.font.size = Pt(font_size)


def add_bullets(document: Document, heading: str, items: list[str]) -> None:
    if not items:
        return
    document.add_heading(heading, level=2)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_docx(plan: dict[str, Any], output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = document.add_heading(str(plan["title"]), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if plan.get("subtitle"):
        subtitle = document.add_paragraph(str(plan["subtitle"]))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {plan['date']}")
    if plan.get("client"):
        client = plan["client"]
        if isinstance(client, dict) and client.get("name"):
            meta.add_run(f" | Client: {client['name']}")

    for paragraph in list_value(plan, "overview"):
        document.add_paragraph(paragraph)

    add_bullets(document, "Safety notes", list_value(plan, "safety_notes"))
    add_bullets(document, "Assumptions", list_value(plan, "assumptions"))

    schedule = plan.get("schedule") or []
    if schedule:
        document.add_heading("Weekly schedule", level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for idx, header in enumerate(["Week", "Day", "Focus", "Duration"]):
            set_cell_text(table.rows[0].cells[idx], header, font_size=8)
        for item in schedule:
            cells = table.add_row().cells
            set_cell_text(cells[0], item.get("week", ""))
            set_cell_text(cells[1], item.get("day", ""))
            set_cell_text(cells[2], item.get("focus", ""))
            set_cell_text(cells[3], item.get("duration", ""))

    document.add_heading("Workout plan", level=2)
    table = document.add_table(rows=1, cols=len(TABLE_COLUMNS))
    table.style = "Table Grid"
    for idx, (_, header) in enumerate(TABLE_COLUMNS):
        set_cell_text(table.rows[0].cells[idx], header, font_size=7)

    for row in iter_exercise_rows(plan):
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(TABLE_COLUMNS):
            value = row.get(key, "")
            if key == "visual_source" and str(value).strip():
                cells[idx].text = ""
                paragraph = cells[idx].paragraphs[0]
                add_hyperlink(paragraph, "source", str(value))
            else:
                set_cell_text(cells[idx], value)

    add_bullets(document, "Progression rules", list_value(plan, "progression_rules"))
    add_bullets(document, "Substitution rules", list_value(plan, "substitution_rules"))
    add_bullets(document, "Cooldown and mobility", list_value(plan, "cooldown"))
    add_bullets(document, "Check-in and logging", list_value(plan, "check_in"))

    warnings = source_warnings(plan)
    if warnings:
        document.add_heading("Source warnings", level=2)
        for warning in warnings:
            document.add_paragraph(warning, style="List Bullet")

    document.save(output)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a workout program DOCX from JSON.")
    parser.add_argument("plan_json", help="Path to workout plan JSON.")
    parser.add_argument("--output", help="Output DOCX path. Defaults to programms/<date>-<slug>.docx.")
    parser.add_argument("--output-dir", default="programms", help="Directory used when --output is omitted.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    plan = load_plan(args.plan_json)
    default_docx, _ = default_output_paths(plan, args.output_dir)
    output = Path(args.output) if args.output else default_docx
    path = build_docx(plan, output)

    for warning in source_warnings(plan):
        print(f"warning: {warning}", file=sys.stderr)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
